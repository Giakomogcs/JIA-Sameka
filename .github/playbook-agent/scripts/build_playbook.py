#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Playbook Agent — builder.

Lê clients/<slug>/content.json e gera clients/<slug>/Playbook.docx com a marca
do cliente. Aborta o build se detectar segredos não sanitizados.

Uso:
    python build_playbook.py --client sameka
    (ou defina a variavel de ambiente PLAYBOOK_CLIENT)

Blocos suportados em content.json:
    paragraph | callout | table | code | steps | screenshot | diagram | bullets
"""
import argparse
import json
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Paleta (couro premium) ----
ACCENT = "B45A3C"
ACCENT2 = "7A3420"
INK = "1A120B"
INK2 = "2B2018"
GRAY = "6B7280"
RULE = "D8CFC7"
ZEBRA = "FAF6F2"
CELLHDR = "F3E5DC"
CODE_BG = "1A120B"
CODE_FG = "EDE7E1"
SHOT_BG = "F7F2ED"
CALLOUT = {
    "info": ("1D4ED8", "EAF0FE", "i"),
    "warning": ("B45309", "FBF1E3", "!"),
    "danger": ("B91C1C", "FBE9E9", "x"),
}
BODY_FONT = "Calibri"
MONO_FONT = "Consolas"

# ---- Guard de sanitização: aborta se achar segredo real ----
FORBIDDEN = [
    (re.compile(r"eyJ[A-Za-z0-9_\-]{20,}"), "JWT real (use <SUPABASE_ANON_KEY>)"),
    (re.compile(r"sk-[A-Za-z0-9]{20,}"), "API key sk-... (use placeholder)"),
    (re.compile(r"longflatworm", re.I), "subdomínio de piloto real (use <SUA-INSTANCIA>)"),
    (re.compile(r"https://[a-z0-9]{8,}\.supabase\.co"), "URL Supabase real (use https://<SEU-PROJETO>.supabase.co)"),
    (re.compile(r"\b\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}\b"), "CNPJ real (use <CNPJ>)"),
]


def guard_sanitization(raw: str):
    problems = []
    for rx, msg in FORBIDDEN:
        for m in rx.finditer(raw):
            problems.append(f"  - {msg}: '{m.group(0)[:40]}'")
    if problems:
        print("BUILD ABORTADO — segredos não sanitizados no content.json:")
        print("\n".join(problems))
        sys.exit(1)


def shade(el, hexcolor):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    el.append(sh)


def set_cell_bg(cell, hexcolor):
    shade(cell._tc.get_or_add_tcPr(), hexcolor)


def cell_borders(cell, color=RULE, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)


def bar_below(p, color, sz):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_inline(p, text, color=INK2, size=10.5):
    token = re.compile(r"(\*\*.+?\*\*|`[^`]+?`)")
    for part in token.split(text):
        if not part:
            continue
        r = p.add_run()
        r.font.name = BODY_FONT
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor.from_string(color)
        if part.startswith("**") and part.endswith("**"):
            r.text = part[2:-2]
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r.text = part[1:-1]
            r.font.name = MONO_FONT
            r.font.color.rgb = RGBColor.from_string(ACCENT2)
        else:
            r.text = part


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_inline(p, text)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text.upper())
        r.font.name = BODY_FONT
        r.font.size = Pt(17)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(ACCENT)
        bar_below(p, ACCENT, 18)
    else:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(13)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(INK)
        bar_below(p, RULE, 6)


def add_callout(doc, variant, title, items):
    color, bg, icon = CALLOUT.get(variant, CALLOUT["info"])
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    cell_borders(cell, color=color, sz=4)
    cell.paragraphs[0].text = ""
    ph = cell.paragraphs[0]
    ph.paragraph_format.space_after = Pt(2)
    rt = ph.add_run(f"[{icon}] {title}")
    rt.bold = True
    rt.font.name = BODY_FONT
    rt.font.size = Pt(10.5)
    rt.font.color.rgb = RGBColor.from_string(color)
    for it in items:
        pb = cell.add_paragraph()
        pb.paragraph_format.left_indent = Pt(10)
        pb.paragraph_format.space_after = Pt(1)
        rb = pb.add_run("• ")
        rb.font.color.rgb = RGBColor.from_string(color)
        rb.font.size = Pt(10)
        add_inline(pb, it, color=INK2, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table_block(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], CELLHDR)
        cell_borders(hdr[i])
        pr = hdr[i].paragraphs[0]
        run = pr.add_run(h)
        run.bold = True
        run.font.name = BODY_FONT
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(ACCENT2)
    for ri, row in enumerate(rows):
        cells = tbl.add_row().cells
        for ci, val in enumerate(row):
            if ri % 2 == 1:
                set_cell_bg(cells[ci], ZEBRA)
            cell_borders(cells[ci])
            pr = cells[ci].paragraphs[0]
            add_inline(pr, str(val), color=INK2, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, CODE_BG)
    cell_borders(cell, color=CODE_BG, sz=2)
    cell.paragraphs[0].text = ""
    for i, line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else " ")
        r.font.name = MONO_FONT
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(CODE_FG)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_diagram(doc, lines):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, ZEBRA)
    cell_borders(cell, color=RULE, sz=4)
    cell.paragraphs[0].text = ""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else " ")
        r.font.name = MONO_FONT
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(INK2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_steps(doc, items, shots_dir):
    for idx, st in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        rn = p.add_run(f"{idx}. ")
        rn.bold = True
        rn.font.color.rgb = RGBColor.from_string(ACCENT)
        rn.font.size = Pt(10.5)
        add_inline(p, st.get("action", ""), color=INK, size=10.5)
        for label, key in (("Onde", "where"), ("O que esperar", "expect")):
            if st.get(key):
                pc = doc.add_paragraph()
                pc.paragraph_format.left_indent = Pt(16)
                pc.paragraph_format.space_after = Pt(0)
                rl = pc.add_run(f"{label}: ")
                rl.bold = True
                rl.font.size = Pt(9.5)
                rl.font.color.rgb = RGBColor.from_string(GRAY)
                add_inline(pc, st[key], color=INK2, size=9.5)
        if st.get("screenshot"):
            add_screenshot(doc, st["screenshot"], None, shots_dir)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_screenshot(doc, path, caption, shots_dir):
    img = shots_dir / path
    if img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img), width=Inches(5.6))
    else:
        tbl = doc.add_table(rows=1, cols=1)
        cell = tbl.rows[0].cells[0]
        set_cell_bg(cell, SHOT_BG)
        cell_borders(cell, color=RULE, sz=4)
        ph = cell.paragraphs[0]
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = ph.add_run(f"[ captura pendente: {path} ]")
        r.font.name = MONO_FONT
        r.font.size = Pt(9)
        r.italic = True
        r.font.color.rgb = RGBColor.from_string(GRAY)
    if caption:
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = pc.add_run(caption)
        rc.italic = True
        rc.font.size = Pt(9)
        rc.font.color.rgb = RGBColor.from_string(GRAY)


def add_cover(doc, data):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(data.get("brand_display", "").upper())
    r.font.name = BODY_FONT
    r.font.size = Pt(40)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bar_below(p2, ACCENT2, 12)
    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = pt.add_run(data.get("title", ""))
    rt.font.size = Pt(20)
    rt.bold = True
    rt.font.color.rgb = RGBColor.from_string(INK)
    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = ps.add_run(data.get("subtitle", ""))
    rs.font.size = Pt(13)
    rs.font.color.rgb = RGBColor.from_string(GRAY)
    for _ in range(2):
        doc.add_paragraph()
    pk = doc.add_paragraph()
    pk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rk = pk.add_run(data.get("stack", ""))
    rk.font.name = MONO_FONT
    rk.font.size = Pt(10)
    rk.font.color.rgb = RGBColor.from_string(ACCENT2)
    pv = doc.add_paragraph()
    pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rv = pv.add_run(f"versão {data.get('version', '1.0')}")
    rv.font.size = Pt(9)
    rv.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_page_break()


def add_footer(doc):
    sec = doc.sections[0]
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Playbook · uso interno · ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r2 = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")
    rpr.append(sz)
    r2.append(rpr)
    fld.append(r2)
    p._p.append(fld)


def render_block(doc, block, shots_dir):
    t = block.get("type")
    if t == "paragraph":
        add_paragraph(doc, block["text"])
    elif t == "callout":
        add_callout(doc, block.get("variant", "info"), block.get("title", ""), block.get("items", []))
    elif t == "table":
        add_table_block(doc, block["headers"], block["rows"])
    elif t == "code":
        add_code(doc, block["text"])
    elif t == "steps":
        add_steps(doc, block["items"], shots_dir)
    elif t == "screenshot":
        add_screenshot(doc, block["path"], block.get("caption"), shots_dir)
    elif t == "diagram":
        add_diagram(doc, block["lines"])
    elif t == "bullets":
        for it in block["items"]:
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Pt(10)
            p.add_run("• ").font.color.rgb = RGBColor.from_string(ACCENT)
            add_inline(p, it, color=INK2, size=10)


def build(client: str):
    root = Path(__file__).resolve().parent.parent
    cdir = root / "clients" / client
    src = cdir / "content.json"
    shots_dir = cdir / "screenshots"
    out = cdir / "Playbook.docx"
    if not src.exists():
        print(f"content.json não encontrado: {src}")
        sys.exit(1)

    raw = src.read_text(encoding="utf-8")
    guard_sanitization(raw)
    data = json.loads(raw)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)

    add_cover(doc, data)
    add_footer(doc)

    for ci, ch in enumerate(data["chapters"]):
        num = ch["id"].split("-")[0]
        add_heading(doc, f"{num} · {ch['title']}", 1)
        for sec in ch.get("sections", []):
            if sec.get("heading"):
                add_heading(doc, sec["heading"], 2)
            for block in sec.get("blocks", []):
                render_block(doc, block, shots_dir)
        if ci < len(data["chapters"]) - 1:
            doc.add_page_break()

    doc.save(str(out))
    size = out.stat().st_size
    n_ch = len(data["chapters"])
    n_steps = sum(
        len(b.get("items", []))
        for ch in data["chapters"]
        for sec in ch.get("sections", [])
        for b in sec.get("blocks", [])
        if b.get("type") == "steps"
    )
    n_shots = sum(
        1
        for ch in data["chapters"]
        for sec in ch.get("sections", [])
        for b in sec.get("blocks", [])
        for _ in ([b] if b.get("type") == "screenshot" else [])
    ) + sum(
        1
        for ch in data["chapters"]
        for sec in ch.get("sections", [])
        for b in sec.get("blocks", [])
        if b.get("type") == "steps"
        for st in b.get("items", [])
        if st.get("screenshot")
    )
    print(f"OK -> {out}")
    print(f"DOCX OK - {size:,} bytes | capítulos={n_ch} passos={n_steps} prints_ref={n_shots}")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--client", default=os.environ.get("PLAYBOOK_CLIENT", ""))
    args = ap.parse_args()
    if not args.client:
        print("Informe --client <slug> ou defina PLAYBOOK_CLIENT.")
        sys.exit(1)
    build(args.client)
