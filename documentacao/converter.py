#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Converte documentacao/DOCUMENTACAO.md em um Word (.docx) estilizado com a marca Sameka.

Parser de Markdown pragmático (não cobre 100% do CommonMark) com suporte a:
  headings (#..####), tabelas (pipe), code fences (``` incl. mermaid),
  callouts > [!NOTE|TIP|WARNING|DANGER], bullets (-), listas numeradas (1.),
  inline **negrito** *itálico* `código`, capa, header/footer, TOC, regra ---.
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Paleta Sameka (couro premium) ----
ACCENT   = "B45A3C"   # terracota couro — H1, capa, header de tabela
ACCENT2  = "7A3420"   # marrom escuro — detalhes, rodapé capa
INK      = "1A120B"   # títulos escuros (H2/H3)
INK2     = "2B2018"   # texto corpo
GRAY     = "6B7280"   # legendas, metadados
RULE     = "D8CFC7"   # bordas de tabela
ZEBRA    = "FAF6F2"   # linhas alternadas
CELLHDR  = "F3E5DC"   # wash do cabeçalho de tabela
CODE_BG  = "1A120B"
CODE_FG  = "EDE7E1"
CALLOUT = {
    "NOTE":    ("1D4ED8", "EAF0FE", "ℹ"),
    "TIP":     ("047857", "E6F4EF", "✔"),
    "WARNING": ("B45309", "FBF1E3", "⚠"),
    "DANGER":  ("B91C1C", "FBE9E9", "✖"),
}
ROOT = Path(__file__).resolve().parent
SRC  = ROOT / "DOCUMENTACAO.md"
OUT  = ROOT / "Sameka-Documentacao.docx"
BODY_FONT = "Calibri"
MONO_FONT = "Consolas"


def shade(el, hexcolor):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    el.append(sh)


def set_cell_bg(cell, hexcolor):
    shade(cell._tc.get_or_add_tcPr(), hexcolor)


def cell_borders(cell, color=RULE, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tcPr.append(borders)


def add_runs_inline(paragraph, text, base_color=INK2, base_size=10.5, bold=False):
    """Renderiza **negrito**, *itálico*, `código` dentro de um parágrafo."""
    token = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")
    for part in token.split(text):
        if not part:
            continue
        r = paragraph.add_run()
        r.font.name = BODY_FONT
        r.font.size = Pt(base_size)
        r.font.color.rgb = RGBColor.from_string(base_color)
        r.bold = bold
        if part.startswith("**") and part.endswith("**"):
            r.text = part[2:-2]
            r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r.text = part[1:-1]
            r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r.text = part[1:-1]
            r.font.name = MONO_FONT
            r.font.color.rgb = RGBColor.from_string(ACCENT2)
        else:
            r.text = part


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text.upper())
        r.font.name = BODY_FONT
        r.font.size = Pt(17)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(ACCENT)
        bar_below(p, ACCENT, 18)
    elif level == 2:
        p.paragraph_format.space_before = Pt(13)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(14)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(INK)
        bar_below(p, RULE, 6)
    elif level == 3:
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(12)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(ACCENT2)
    else:
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(INK2)
    return p


def bar_below(paragraph, color, sz):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_runs_inline(p, text)
    return p


def add_bullet(doc, text, numbered=False, num=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    prefix = f"{num}. " if numbered else "•  "
    r = p.add_run(prefix)
    r.font.name = BODY_FONT
    r.font.size = Pt(10.5)
    r.bold = numbered
    r.font.color.rgb = RGBColor.from_string(ACCENT if numbered else ACCENT2)
    add_runs_inline(p, text)
    return p


def add_code_block(doc, lines, caption=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, CODE_BG)
    cell_borders(cell, CODE_BG, 4)
    cell.paragraphs[0].text = ""
    for i, ln in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(ln if ln else " ")
        r.font.name = MONO_FONT
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(CODE_FG)
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = c.add_run(caption)
        rc.italic = True
        rc.font.size = Pt(8.5)
        rc.font.name = BODY_FONT
        rc.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc, kind, lines):
    color, bg, icon = CALLOUT.get(kind, CALLOUT["NOTE"])
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    # barra lateral grossa colorida (left border)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "0"); left.set(qn("w:color"), color)
    borders.append(left)
    for edge in ("top", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), bg)
        borders.append(e)
    tcPr.append(borders)
    first = cell.paragraphs[0]
    first.paragraph_format.space_after = Pt(2)
    rh = first.add_run(f"{icon}  {kind}")
    rh.bold = True
    rh.font.name = BODY_FONT
    rh.font.size = Pt(10.5)
    rh.font.color.rgb = RGBColor.from_string(color)
    for ln in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        add_runs_inline(p, ln, base_color=INK2, base_size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(header):
        set_cell_bg(hdr[i], ACCENT)
        cell_borders(hdr[i])
        para = hdr[i].paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        add_runs_inline(para, htext, base_color="FFFFFF", base_size=10, bold=True)
        for r in para.runs:
            r.bold = True
            r.font.color.rgb = RGBColor.from_string("FFFFFF")
    for ri, row in enumerate(body):
        cells = tbl.add_row().cells
        for ci, ctext in enumerate(row):
            if ci < len(cells):
                if ri % 2 == 1:
                    set_cell_bg(cells[ci], ZEBRA)
                cell_borders(cells[ci])
                para = cells[ci].paragraphs[0]
                para.paragraph_format.space_after = Pt(0)
                add_runs_inline(para, ctext, base_size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    bar_below(p, ACCENT, 40)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run("SAMEKA")
    r.font.name = BODY_FONT
    r.bold = True
    r.font.size = Pt(46)
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    sub = doc.add_paragraph()
    rs = sub.add_run("Documentação Técnica e de Negócio")
    rs.font.name = BODY_FONT
    rs.font.size = Pt(20)
    rs.font.color.rgb = RGBColor.from_string(INK)
    desc = doc.add_paragraph()
    rd = desc.add_run("Copiloto Estratégico de Vendas — calçados de couro premium para bebê")
    rd.font.name = BODY_FONT
    rd.italic = True
    rd.font.size = Pt(12)
    rd.font.color.rgb = RGBColor.from_string(GRAY)
    p2 = doc.add_paragraph()
    bar_below(p2, ACCENT2, 18)
    for _ in range(8):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    rm = meta.add_run("Projeto: Sameka  ·  Data: 2026-06-23  ·  Autor: Doc Master")
    rm.font.name = BODY_FONT
    rm.font.size = Pt(10.5)
    rm.font.color.rgb = RGBColor.from_string(ACCENT2)
    doc.add_page_break()


def add_toc(doc):
    h = doc.add_paragraph()
    r = h.add_run("Sumário")
    r.font.name = BODY_FONT
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    bar_below(h, RULE, 6)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Atualize o sumário (F9 / clique direito → Atualizar campo)."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(t)
    run._r.append(fld_end)
    doc.add_page_break()


def add_page_number(paragraph):
    run = paragraph.add_run("Página ")
    run.font.size = Pt(8); run.font.color.rgb = RGBColor.from_string(GRAY)
    for instr_text in ("PAGE",):
        b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
        i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = instr_text
        e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
        run._r.append(b); run._r.append(i); run._r.append(e)
    run2 = paragraph.add_run(" de ")
    run2.font.size = Pt(8); run2.font.color.rgb = RGBColor.from_string(GRAY)
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = "NUMPAGES"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run2._r.append(b); run2._r.append(i); run2._r.append(e)


def setup_header_footer(doc):
    sec = doc.sections[0]
    header = sec.header
    hp = header.paragraphs[0]
    hp.text = ""
    r = hp.add_run("Sameka — Documentação")
    r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(GRAY); r.font.name = BODY_FONT
    tab = hp.add_run("\tProjeto Sameka")
    tab.font.size = Pt(8); tab.font.color.rgb = RGBColor.from_string(GRAY)
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)


def render(md_lines, doc):
    i = 0
    n = len(md_lines)
    while i < n:
        line = md_lines[i].rstrip("\n")
        stripped = line.strip()

        # code fence
        m = re.match(r"^```(\w*)", stripped)
        if m:
            lang = m.group(1)
            block = []
            i += 1
            while i < n and not md_lines[i].strip().startswith("```"):
                block.append(md_lines[i].rstrip("\n"))
                i += 1
            i += 1  # skip closing fence
            cap = "(diagrama Mermaid)" if lang == "mermaid" else None
            add_code_block(doc, block, caption=cap)
            continue

        # callout
        mc = re.match(r"^>\s*\[!(\w+)\]", stripped)
        if mc:
            kind = mc.group(1).upper()
            clines = []
            i += 1
            while i < n and md_lines[i].lstrip().startswith(">"):
                txt = md_lines[i].lstrip()[1:].strip()
                if txt:
                    clines.append(txt)
                i += 1
            add_callout(doc, kind, clines)
            continue

        # table
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|", md_lines[i + 1].strip()):
            rows = [parse_table_row(stripped)]
            i += 2  # skip header + separator
            while i < n and md_lines[i].strip().startswith("|"):
                rows.append(parse_table_row(md_lines[i].strip()))
                i += 1
            add_table(doc, rows)
            continue

        # headings
        mh = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if mh:
            add_heading(doc, mh.group(2).strip(), len(mh.group(1)))
            i += 1
            continue

        # horizontal rule
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # blockquote (simple)
        if stripped.startswith(">"):
            add_paragraph(doc, stripped.lstrip(">").strip())
            i += 1
            continue

        # numbered list
        mn = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if mn:
            add_bullet(doc, mn.group(2), numbered=True, num=mn.group(1))
            i += 1
            continue

        # bullet
        if stripped.startswith("- ") or stripped.startswith("* "):
            add_bullet(doc, stripped[2:])
            i += 1
            continue

        # blank
        if not stripped:
            i += 1
            continue

        add_paragraph(doc, stripped)
        i += 1


def main():
    md = SRC.read_text(encoding="utf-8").splitlines(keepends=True)
    # remove a primeira H1 e o blockquote inicial (vão para a capa)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor.from_string(INK2)

    add_cover(doc)
    add_toc(doc)
    setup_header_footer(doc)

    # pula o título principal (linha 1) e a citação inicial até o primeiro '---'
    start = 0
    for idx, ln in enumerate(md):
        if ln.strip() == "---":
            start = idx + 1
            break
    render(md[start:], doc)

    doc.save(OUT)
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    main()
